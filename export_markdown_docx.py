#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int) -> None:
    level = max(1, min(level, 4))
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def add_table_from_markdown(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    # skip the markdown separator row
    header = rows[0]
    body = [r for r in rows[2:] if len(r) == len(header)]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for j, t in enumerate(header):
        table.cell(0, j).text = t
    for i, row in enumerate(body, start=1):
        for j, t in enumerate(row):
            table.cell(i, j).text = t


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_style(doc)

    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, code_buf)
                code_buf = []
            continue

        if in_code:
            code_buf.append(line)
            continue

        if stripped.startswith("|"):
            table_buf.append(line)
            continue
        elif table_buf:
            add_table_from_markdown(doc, table_buf)
            table_buf = []

        if not stripped:
            doc.add_paragraph("")
            continue

        img_match = re.match(r"!\[[^\]]*]\(([^)]+)\)", stripped)
        if img_match:
            img_rel = img_match.group(1).strip()
            img_path = (md_path.parent / img_rel).resolve()
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.2))
                except Exception:
                    doc.add_paragraph(f"[图片插入失败] {img_rel}")
            else:
                doc.add_paragraph(f"[图片不存在] {img_rel}")
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            add_heading(doc, text, level)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(stripped, style="List Number")
            continue

        doc.add_paragraph(stripped)

    if table_buf:
        add_table_from_markdown(doc, table_buf)
    if code_buf:
        add_code_block(doc, code_buf)

    doc.save(str(docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Export markdown report to DOCX.")
    parser.add_argument(
        "--input",
        type=str,
        default="REPORT_MAIN_PDF_STYLE.md",
        help="Input markdown path",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="REPORT_MAIN_PDF_STYLE.docx",
        help="Output docx path",
    )
    args = parser.parse_args()

    md_path = Path(args.input).resolve()
    docx_path = Path(args.output).resolve()
    if not md_path.exists():
        raise FileNotFoundError(f"Input markdown not found: {md_path}")

    convert(md_path, docx_path)
    print(f"DOCX generated: {docx_path}")


if __name__ == "__main__":
    main()
